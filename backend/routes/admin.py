from flask import Blueprint, request, jsonify, send_file
import mysql.connector
from mysql.connector import Error
from decimal import Decimal
import bcrypt
from io import BytesIO
from datetime import datetime
import os
import uuid
from werkzeug.utils import secure_filename

admin_bp = Blueprint('admin', __name__)

# Configuration
BACKEND_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
UPLOAD_DIR = os.path.join(BACKEND_DIR, 'static', 'uploads', 'products')

def get_db():
    """Get database connection"""
    try:
        DB_CONFIG = {
            'host': 'localhost',
            'user': 'root',
            'password': '', 
            'database': 'ecommerce_db'
        }
        return mysql.connector.connect(**DB_CONFIG)
    except Error as e:
        print(f"❌ Database error: {e}")
        return None

def get_date_grouping(period):
    """
    Get SQL date grouping and formatting based on period.
    Returns a tuple of (date_group, date_format).
    
    This function uses a whitelist approach to prevent SQL injection.
    Only pre-defined, hardcoded values are returned.
    """
    # Whitelist of allowed period values with their corresponding SQL expressions
    period_config = {
        'daily': ("DATE(o.created_at)", "%Y-%m-%d"),
        'weekly': ("YEARWEEK(o.created_at, 1)", "Week %v, %Y"),
        'monthly': ("DATE_FORMAT(o.created_at, '%Y-%m')", "%Y-%m")
    }
    
    # If period is not in whitelist, default to daily
    return period_config.get(period, period_config['daily'])

# ========== ADMIN ORDER ENDPOINTS ==========
@admin_bp.route('/orders', methods=['GET'])
def get_all_orders():
    """Get all orders (admin view)"""
    try:
        conn = get_db()
        if not conn:
            return jsonify({'error': 'Database connection failed'}), 500
        
        cursor = conn.cursor(dictionary=True)
        
        # FIXED: Calculate total quantity, not just row count
        cursor.execute("""
            SELECT 
                o.*, 
                u.first_name, 
                u.last_name, 
                u.email,
                (SELECT SUM(quantity) FROM order_items WHERE order_id = o.id) as total_quantity,
                (SELECT COUNT(*) FROM order_items WHERE order_id = o.id) as item_count
            FROM orders o
            LEFT JOIN users u ON o.user_id = u.id
            ORDER BY o.created_at DESC
        """)
        
        orders = cursor.fetchall()
        
        # Get order items for each order
        for order in orders:
            cursor.execute("""
                SELECT 
                    oi.*, 
                    p.name, 
                    p.image_url,
                    p.description,
                    oi.quantity  -- Make sure to get quantity from order_items
                FROM order_items oi
                JOIN products p ON oi.product_id = p.id
                WHERE oi.order_id = %s
            """, (order['id'],))
            
            items = cursor.fetchall()
            
            # Convert Decimal to float
            for item in items:
                if 'price_at_time' in item and isinstance(item['price_at_time'], Decimal):
                    item['price_at_time'] = float(item['price_at_time'])
                # Ensure quantity is integer
                if 'quantity' in item:
                    item['quantity'] = int(item['quantity'])
            
            order['items'] = items
        
        cursor.close()
        conn.close()
        
        # Convert Decimal to float and ensure integer counts
        for order in orders:
            if 'total_amount' in order and isinstance(order['total_amount'], Decimal):
                order['total_amount'] = float(order['total_amount'])
            if 'total_quantity' in order:
                order['total_quantity'] = int(order['total_quantity']) if order['total_quantity'] else 0
            if 'item_count' in order:
                order['item_count'] = int(order['item_count']) if order['item_count'] else 0
        
        return jsonify(orders), 200
        
    except Exception as e:
        print(f"❌ Get all orders error: {e}")
        return jsonify({'error': str(e)}), 500

@admin_bp.route('/orders/<int:order_id>/approve', methods=['PUT'])
def approve_order(order_id):
    """Approve order (admin action)"""
    try:
        conn = get_db()
        if not conn:
            return jsonify({'error': 'Database connection failed'}), 500
        
        cursor = conn.cursor(dictionary=True)
        
        # Check if order exists and is pending
        cursor.execute("SELECT status FROM orders WHERE id = %s", (order_id,))
        order = cursor.fetchone()
        
        if not order:
            cursor.close()
            conn.close()
            return jsonify({'error': 'Order not found'}), 404
        
        if order['status'] != 'pending':
            cursor.close()
            conn.close()
            return jsonify({'error': 'Only pending orders can be approved'}), 400
        
        # Update order status to processing
        cursor.execute("UPDATE orders SET status = 'processing' WHERE id = %s", (order_id,))
        
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({'message': 'Order approved successfully'}), 200
        
    except Exception as e:
        print(f"❌ Approve order error: {e}")
        return jsonify({'error': str(e)}), 500

@admin_bp.route('/orders/<int:order_id>/decline', methods=['PUT'])
def decline_order(order_id):
    """Decline order with reason (admin action)"""
    try:
        data = request.get_json()
        decline_reason = data.get('decline_reason')
        
        if not decline_reason or not decline_reason.strip():
            return jsonify({'error': 'Decline reason is required'}), 400
        
        conn = get_db()
        if not conn:
            return jsonify({'error': 'Database connection failed'}), 500
        
        cursor = conn.cursor(dictionary=True)
        
        # Check if order exists and is pending
        cursor.execute("SELECT status FROM orders WHERE id = %s", (order_id,))
        order = cursor.fetchone()
        
        if not order:
            cursor.close()
            conn.close()
            return jsonify({'error': 'Order not found'}), 404
        
        if order['status'] != 'pending':
            cursor.close()
            conn.close()
            return jsonify({'error': 'Only pending orders can be declined'}), 400
        
        # Get order items to replenish stock
        cursor.execute("""
            SELECT product_id, quantity 
            FROM order_items 
            WHERE order_id = %s
        """, (order_id,))
        order_items = cursor.fetchall()
        
        # Replenish stock for each item
        for item in order_items:
            cursor.execute("""
                UPDATE products 
                SET stock_quantity = stock_quantity + %s 
                WHERE id = %s
            """, (item['quantity'], item['product_id']))
        
        # Update order status to declined with reason
        cursor.execute("""
            UPDATE orders 
            SET status = 'declined', decline_reason = %s 
            WHERE id = %s
        """, (decline_reason, order_id))
        
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({'message': 'Order declined successfully'}), 200
        
    except Exception as e:
        print(f"❌ Decline order error: {e}")
        return jsonify({'error': str(e)}), 500

# ========== ADMIN PRODUCT MANAGEMENT ENDPOINTS ==========
@admin_bp.route('/products', methods=['POST'])
def create_product():
    """Create a new product (admin action)"""
    try:
        data = request.get_json()
        
        required_fields = ['name', 'description', 'price', 'category_id', 'stock_quantity']
        for field in required_fields:
            if field not in data:
                return jsonify({'error': f'{field} is required'}), 400
        
        conn = get_db()
        if not conn:
            return jsonify({'error': 'Database connection failed'}), 500
        
        cursor = conn.cursor(dictionary=True)
        
        cursor.execute("""
            INSERT INTO products (name, description, price, category_id, stock_quantity, image_url)
            VALUES (%s, %s, %s, %s, %s, %s)
        """, (
            data['name'],
            data['description'],
            data['price'],
            data['category_id'],
            data['stock_quantity'],
            data.get('image_url', '')
        ))
        
        product_id = cursor.lastrowid
        
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({
            'message': 'Product created successfully',
            'product_id': product_id
        }), 201
        
    except Exception as e:
        print(f"❌ Create product error: {e}")
        return jsonify({'error': str(e)}), 500

@admin_bp.route('/products/<int:product_id>', methods=['PUT'])
def update_product(product_id):
    """Update a product (admin action)"""
    try:
        data = request.get_json()
        
        conn = get_db()
        if not conn:
            return jsonify({'error': 'Database connection failed'}), 500
        
        cursor = conn.cursor(dictionary=True)
        
        # Build update query dynamically based on provided fields
        update_fields = []
        params = []
        
        if 'name' in data:
            update_fields.append('name = %s')
            params.append(data['name'])
        if 'description' in data:
            update_fields.append('description = %s')
            params.append(data['description'])
        if 'price' in data:
            update_fields.append('price = %s')
            params.append(data['price'])
        if 'category_id' in data:
            update_fields.append('category_id = %s')
            params.append(data['category_id'])
        if 'stock_quantity' in data:
            update_fields.append('stock_quantity = %s')
            params.append(data['stock_quantity'])
        if 'image_url' in data:
            update_fields.append('image_url = %s')
            params.append(data['image_url'])
        if 'is_active' in data:
            update_fields.append('is_active = %s')
            params.append(data['is_active'])
        
        if not update_fields:
            cursor.close()
            conn.close()
            return jsonify({'error': 'No fields to update'}), 400
        
        params.append(product_id)
        
        # Dito Safe to use f-string here ng update_fields list is built from a validated whitelist of field names above.
        query = f"UPDATE products SET {', '.join(update_fields)} WHERE id = %s"
        
        cursor.execute(query, params)
        
        if cursor.rowcount == 0:
            cursor.close()
            conn.close()
            return jsonify({'error': 'Product not found'}), 404
        
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({'message': 'Product updated successfully'}), 200
        
    except Exception as e:
        print(f"❌ Update product error: {e}")
        return jsonify({'error': str(e)}), 500

@admin_bp.route('/products/<int:product_id>', methods=['DELETE'])
def delete_product(product_id):
    """Delete a product (admin action) - soft delete by setting is_active to false"""
    try:
        conn = get_db()
        if not conn:
            return jsonify({'error': 'Database connection failed'}), 500
        
        cursor = conn.cursor()
        
        # Soft delete - set is_active to false
        cursor.execute("UPDATE products SET is_active = FALSE WHERE id = %s", (product_id,))
        
        if cursor.rowcount == 0:
            cursor.close()
            conn.close()
            return jsonify({'error': 'Product not found'}), 404
        
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({'message': 'Product deleted successfully'}), 200
        
    except Exception as e:
        print(f"❌ Delete product error: {e}")
        return jsonify({'error': str(e)}), 500

# ========== ADMIN CATEGORY MANAGEMENT ENDPOINTS ==========
@admin_bp.route('/categories', methods=['POST'])
def create_category():
    """Create a new category (admin action)"""
    try:
        data = request.get_json()
        
        if 'name' not in data or not data['name'].strip():
            return jsonify({'error': 'Category name is required'}), 400
        
        conn = get_db()
        if not conn:
            return jsonify({'error': 'Database connection failed'}), 500
        
        cursor = conn.cursor(dictionary=True)
        
        cursor.execute("""
            INSERT INTO categories (name, description)
            VALUES (%s, %s)
        """, (data['name'], data.get('description', '')))
        
        category_id = cursor.lastrowid
        
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({
            'message': 'Category created successfully',
            'category_id': category_id
        }), 201
        
    except Exception as e:
        print(f"❌ Create category error: {e}")
        return jsonify({'error': str(e)}), 500

@admin_bp.route('/categories/<int:category_id>', methods=['PUT'])
def update_category(category_id):
    """Update a category (admin action)"""
    try:
        data = request.get_json()
        
        conn = get_db()
        if not conn:
            return jsonify({'error': 'Database connection failed'}), 500
        
        cursor = conn.cursor()
        
        # Build update query
        update_fields = []
        params = []
        
        if 'name' in data:
            update_fields.append('name = %s')
            params.append(data['name'])
        if 'description' in data:
            update_fields.append('description = %s')
            params.append(data['description'])
        
        if not update_fields:
            cursor.close()
            conn.close()
            return jsonify({'error': 'No fields to update'}), 400
        
        params.append(category_id)
        # Dito Safe to use f-string here ng update_fields list is built from a validated whitelist of field names above.
        query = f"UPDATE categories SET {', '.join(update_fields)} WHERE id = %s"
        
        cursor.execute(query, params)
        
        if cursor.rowcount == 0:
            cursor.close()
            conn.close()
            return jsonify({'error': 'Category not found'}), 404
        
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({'message': 'Category updated successfully'}), 200
        
    except Exception as e:
        print(f"❌ Update category error: {e}")
        return jsonify({'error': str(e)}), 500

@admin_bp.route('/categories/<int:category_id>', methods=['DELETE'])
def delete_category(category_id):
    """Delete a category (admin action)"""
    try:
        conn = get_db()
        if not conn:
            return jsonify({'error': 'Database connection failed'}), 500
        
        cursor = conn.cursor()
        
        # Check if category has products
        cursor.execute("SELECT COUNT(*) as count FROM products WHERE category_id = %s", (category_id,))
        result = cursor.fetchone()
        
        if result[0] > 0:
            cursor.close()
            conn.close()
            return jsonify({'error': 'Cannot delete category with existing products'}), 400
        
        cursor.execute("DELETE FROM categories WHERE id = %s", (category_id,))
        
        if cursor.rowcount == 0:
            cursor.close()
            conn.close()
            return jsonify({'error': 'Category not found'}), 404
        
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({'message': 'Category deleted successfully'}), 200
        
    except Exception as e:
        print(f"❌ Delete category error: {e}")
        return jsonify({'error': str(e)}), 500

# ========== ADMIN SALES REPORTS ENDPOINTS ==========
@admin_bp.route('/reports/sales', methods=['GET'])
def get_sales_report():
    """Get sales reports with daily, weekly, or monthly aggregation"""
    try:
        period = request.args.get('period', 'daily')  # daily, weekly, monthly
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')
        
        if period not in ['daily', 'weekly', 'monthly']:
            return jsonify({'error': 'Invalid period. Use daily, weekly, or monthly'}), 400
        
        conn = get_db()
        if not conn:
            return jsonify({'error': 'Database connection failed'}), 500
        
        cursor = conn.cursor(dictionary=True)
        
        # Build date filter
        date_filter = ""
        params = []
        
        if start_date:
            date_filter += " AND DATE(o.created_at) >= %s"
            params.append(start_date)
        if end_date:
            date_filter += " AND DATE(o.created_at) <= %s"
            params.append(end_date)
        
        # Get date grouping configuration from helper function (prevents SQL injection)
        date_group, date_format = get_date_grouping(period)
        
        # Dito rin Safe to use f-string here: date_group and date_format come from a validated whitelist
        query = f"""
            SELECT 
                {date_group} as period,
                DATE_FORMAT(o.created_at, '{date_format}') as period_label,
                COUNT(DISTINCT o.id) as total_orders,
                SUM(o.total_amount) as total_revenue,
                AVG(o.total_amount) as average_order_value,
                COUNT(DISTINCT o.user_id) as unique_customers,
                SUM(CASE WHEN o.status = 'delivered' THEN 1 ELSE 0 END) as completed_orders,
                SUM(CASE WHEN o.status = 'cancelled' THEN 1 ELSE 0 END) as cancelled_orders,
                SUM(CASE WHEN o.status = 'declined' THEN 1 ELSE 0 END) as declined_orders
            FROM orders o
            WHERE o.status IN ('pending', 'processing', 'shipped', 'in_transit', 'delivered', 'cancelled', 'declined')
            {date_filter}
            GROUP BY {date_group}
            ORDER BY period DESC
            LIMIT 50
        """
        
        cursor.execute(query, params)
        sales_data = cursor.fetchall()
        
        # Get overall statistics
        stats_query = f"""
            SELECT 
                COUNT(DISTINCT o.id) as total_orders,
                SUM(o.total_amount) as total_revenue,
                AVG(o.total_amount) as average_order_value,
                COUNT(DISTINCT o.user_id) as unique_customers,
                SUM(CASE WHEN o.status = 'delivered' THEN o.total_amount ELSE 0 END) as delivered_revenue,
                COUNT(CASE WHEN o.status = 'delivered' THEN 1 END) as delivered_orders
            FROM orders o
            WHERE 1=1 {date_filter}
        """
        
        cursor.execute(stats_query, params)
        overall_stats = cursor.fetchone()
        
        # Get top selling products
        top_products_query = f"""
            SELECT 
                p.id,
                p.name,
                p.image_url,
                SUM(oi.quantity) as total_quantity_sold,
                SUM(oi.quantity * oi.price_at_time) as total_revenue
            FROM order_items oi
            JOIN products p ON oi.product_id = p.id
            JOIN orders o ON oi.order_id = o.id
            WHERE o.status IN ('delivered', 'shipped', 'in_transit', 'processing')
            {date_filter}
            GROUP BY p.id, p.name, p.image_url
            ORDER BY total_quantity_sold DESC
            LIMIT 10
        """
        
        cursor.execute(top_products_query, params)
        top_products = cursor.fetchall()
        
        cursor.close()
        conn.close()
        
        # Convert Decimal to float
        for item in sales_data:
            if 'total_revenue' in item and item['total_revenue'] is not None:
                item['total_revenue'] = float(item['total_revenue'])
            if 'average_order_value' in item and item['average_order_value'] is not None:
                item['average_order_value'] = float(item['average_order_value'])
        
        if overall_stats:
            if 'total_revenue' in overall_stats and overall_stats['total_revenue'] is not None:
                overall_stats['total_revenue'] = float(overall_stats['total_revenue'])
            if 'average_order_value' in overall_stats and overall_stats['average_order_value'] is not None:
                overall_stats['average_order_value'] = float(overall_stats['average_order_value'])
            if 'delivered_revenue' in overall_stats and overall_stats['delivered_revenue'] is not None:
                overall_stats['delivered_revenue'] = float(overall_stats['delivered_revenue'])
        
        for product in top_products:
            if 'total_revenue' in product and product['total_revenue'] is not None:
                product['total_revenue'] = float(product['total_revenue'])
        
        return jsonify({
            'period': period,
            'start_date': start_date,
            'end_date': end_date,
            'sales_data': sales_data,
            'overall_stats': overall_stats,
            'top_products': top_products
        }), 200
        
    except Exception as e:
        print(f"❌ Get sales report error: {e}")
        return jsonify({'error': str(e)}), 500

@admin_bp.route('/reports/sales/generate', methods=['POST'])
def generate_sales_report():
    """Generate PDF or DOCX sales report"""
    try:
        data = request.get_json()
        format_type = data.get('format', 'pdf')  # pdf or docx
        period = data.get('period', 'daily')
        start_date = data.get('start_date')
        end_date = data.get('end_date')
        
        if format_type not in ['pdf', 'docx']:
            return jsonify({'error': 'Invalid format. Use pdf or docx'}), 400
        
        if period not in ['daily', 'weekly', 'monthly']:
            return jsonify({'error': 'Invalid period. Use daily, weekly, or monthly'}), 400
        
        # Get sales data using existing report logic
        conn = get_db()
        if not conn:
            return jsonify({'error': 'Database connection failed'}), 500
        
        cursor = conn.cursor(dictionary=True)
        
        # Build date filter
        date_filter = ""
        params = []
        
        if start_date:
            date_filter += " AND DATE(o.created_at) >= %s"
            params.append(start_date)
        if end_date:
            date_filter += " AND DATE(o.created_at) <= %s"
            params.append(end_date)
        
        # Get date grouping configuration from helper function (prevents SQL injection)
        date_group, date_format = get_date_grouping(period)
        
        # Dito rin Safe to use f-string here: date_group and date_format come from a validated whitelist
        query = f"""
            SELECT 
                {date_group} as period,
                DATE_FORMAT(o.created_at, '{date_format}') as period_label,
                COUNT(DISTINCT o.id) as total_orders,
                SUM(o.total_amount) as total_revenue,
                AVG(o.total_amount) as average_order_value,
                COUNT(DISTINCT o.user_id) as unique_customers,
                SUM(CASE WHEN o.status = 'delivered' THEN 1 ELSE 0 END) as completed_orders,
                SUM(CASE WHEN o.status = 'cancelled' THEN 1 ELSE 0 END) as cancelled_orders,
                SUM(CASE WHEN o.status = 'declined' THEN 1 ELSE 0 END) as declined_orders
            FROM orders o
            WHERE o.status IN ('pending', 'processing', 'shipped', 'in_transit', 'delivered', 'cancelled', 'declined')
            {date_filter}
            GROUP BY {date_group}
            ORDER BY period DESC
            LIMIT 50
        """
        
        cursor.execute(query, params)
        sales_data = cursor.fetchall()
        
        # Get overall statistics
        stats_query = f"""
            SELECT 
                COUNT(DISTINCT o.id) as total_orders,
                SUM(o.total_amount) as total_revenue,
                AVG(o.total_amount) as average_order_value,
                COUNT(DISTINCT o.user_id) as unique_customers,
                SUM(CASE WHEN o.status = 'delivered' THEN o.total_amount ELSE 0 END) as delivered_revenue,
                COUNT(CASE WHEN o.status = 'delivered' THEN 1 END) as delivered_orders
            FROM orders o
            WHERE 1=1 {date_filter}
        """
        
        cursor.execute(stats_query, params)
        overall_stats = cursor.fetchone()
        
        # Get top selling products
        top_products_query = f"""
            SELECT 
                p.id,
                p.name,
                p.image_url,
                SUM(oi.quantity) as total_quantity_sold,
                SUM(oi.quantity * oi.price_at_time) as total_revenue
            FROM order_items oi
            JOIN products p ON oi.product_id = p.id
            JOIN orders o ON oi.order_id = o.id
            WHERE o.status IN ('delivered', 'shipped', 'in_transit', 'processing')
            {date_filter}
            GROUP BY p.id, p.name, p.image_url
            ORDER BY total_quantity_sold DESC
            LIMIT 10
        """
        
        cursor.execute(top_products_query, params)
        top_products = cursor.fetchall()
        
        cursor.close()
        conn.close()
        
        # Convert Decimal to float
        for item in sales_data:
            if 'total_revenue' in item and item['total_revenue'] is not None:
                item['total_revenue'] = float(item['total_revenue'])
            if 'average_order_value' in item and item['average_order_value'] is not None:
                item['average_order_value'] = float(item['average_order_value'])
        
        if overall_stats:
            if 'total_revenue' in overall_stats and overall_stats['total_revenue'] is not None:
                overall_stats['total_revenue'] = float(overall_stats['total_revenue'])
            if 'average_order_value' in overall_stats and overall_stats['average_order_value'] is not None:
                overall_stats['average_order_value'] = float(overall_stats['average_order_value'])
            if 'delivered_revenue' in overall_stats and overall_stats['delivered_revenue'] is not None:
                overall_stats['delivered_revenue'] = float(overall_stats['delivered_revenue'])
        
        for product in top_products:
            if 'total_revenue' in product and product['total_revenue'] is not None:
                product['total_revenue'] = float(product['total_revenue'])
        
        # Generate report based on format
        if format_type == 'pdf':
            file_buffer = generate_pdf_report(period, start_date, end_date, sales_data, overall_stats, top_products)
            filename = f"sales_report_{period}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            mimetype = 'application/pdf'
        else:  # docx
            file_buffer = generate_docx_report(period, start_date, end_date, sales_data, overall_stats, top_products)
            filename = f"sales_report_{period}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        
        return send_file(
            file_buffer,
            mimetype=mimetype,
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as e:
        print(f"❌ Generate sales report error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

def generate_pdf_report(period, start_date, end_date, sales_data, overall_stats, top_products):
    """Generate PDF sales report using reportlab"""
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    elements = []
    styles = getSampleStyleSheet()
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#2c3e50'),
        spaceAfter=30,
        alignment=TA_CENTER
    )
    elements.append(Paragraph("E-Commerce Sales Report", title_style))
    elements.append(Spacer(1, 0.2*inch))
    
    # Report details
    detail_style = styles['Normal']
    elements.append(Paragraph(f"<b>Period:</b> {period.capitalize()}", detail_style))
    if start_date:
        elements.append(Paragraph(f"<b>Start Date:</b> {start_date}", detail_style))
    if end_date:
        elements.append(Paragraph(f"<b>End Date:</b> {end_date}", detail_style))
    elements.append(Paragraph(f"<b>Generated:</b> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", detail_style))
    elements.append(Spacer(1, 0.3*inch))
    
    # Overall Statistics
    elements.append(Paragraph("<b>Overall Statistics</b>", styles['Heading2']))
    elements.append(Spacer(1, 0.1*inch))
    
    if overall_stats:
        stats_data = [
            ['Metric', 'Value'],
            ['Total Orders', str(overall_stats.get('total_orders', 0))],
            ['Total Revenue', f"${overall_stats.get('total_revenue', 0):,.2f}"],
            ['Average Order Value', f"${overall_stats.get('average_order_value', 0):,.2f}"],
            ['Unique Customers', str(overall_stats.get('unique_customers', 0))],
            ['Delivered Orders', str(overall_stats.get('delivered_orders', 0))],
            ['Delivered Revenue', f"${overall_stats.get('delivered_revenue', 0):,.2f}"],
        ]
        
        stats_table = Table(stats_data, colWidths=[3*inch, 2*inch])
        stats_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3498db')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        elements.append(stats_table)
        elements.append(Spacer(1, 0.3*inch))
    
    # Sales Data by Period
    if sales_data:
        elements.append(Paragraph(f"<b>Sales Data by {period.capitalize()}</b>", styles['Heading2']))
        elements.append(Spacer(1, 0.1*inch))
        
        sales_table_data = [['Period', 'Orders', 'Revenue', 'Avg Order', 'Customers']]
        for item in sales_data[:20]:  # Limit to 20 rows for space
            sales_table_data.append([
                str(item.get('period_label', '')),
                str(item.get('total_orders', 0)),
                f"${item.get('total_revenue', 0):,.2f}",
                f"${item.get('average_order_value', 0):,.2f}",
                str(item.get('unique_customers', 0))
            ])
        
        sales_table = Table(sales_table_data, colWidths=[1.5*inch, 1*inch, 1.2*inch, 1.2*inch, 1*inch])
        sales_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3498db')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        elements.append(sales_table)
        elements.append(Spacer(1, 0.3*inch))
    
    # Top Products
    if top_products:
        elements.append(Paragraph("<b>Top 10 Selling Products</b>", styles['Heading2']))
        elements.append(Spacer(1, 0.1*inch))
        
        products_table_data = [['Rank', 'Product Name', 'Quantity Sold', 'Revenue']]
        for idx, product in enumerate(top_products, 1):
            products_table_data.append([
                str(idx),
                str(product.get('name', ''))[:30],  # Truncate long names
                str(product.get('total_quantity_sold', 0)),
                f"${product.get('total_revenue', 0):,.2f}"
            ])
        
        products_table = Table(products_table_data, colWidths=[0.7*inch, 2.5*inch, 1.3*inch, 1.5*inch])
        products_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3498db')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('ALIGN', (0, 0), (0, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        elements.append(products_table)
    
    # Build PDF
    doc.build(elements)
    buffer.seek(0)
    return buffer

def generate_docx_report(period, start_date, end_date, sales_data, overall_stats, top_products):
    """Generate DOCX sales report using python-docx"""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Title
    title = doc.add_heading('E-Commerce Sales Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Report details
    doc.add_paragraph(f"Period: {period.capitalize()}")
    if start_date:
        doc.add_paragraph(f"Start Date: {start_date}")
    if end_date:
        doc.add_paragraph(f"End Date: {end_date}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()
    
    # Overall Statistics
    doc.add_heading('Overall Statistics', level=1)
    if overall_stats:
        stats_table = doc.add_table(rows=7, cols=2)
        stats_table.style = 'Light Grid Accent 1'
        
        stats_table.rows[0].cells[0].text = 'Metric'
        stats_table.rows[0].cells[1].text = 'Value'
        stats_table.rows[1].cells[0].text = 'Total Orders'
        stats_table.rows[1].cells[1].text = str(overall_stats.get('total_orders', 0))
        stats_table.rows[2].cells[0].text = 'Total Revenue'
        stats_table.rows[2].cells[1].text = f"${overall_stats.get('total_revenue', 0):,.2f}"
        stats_table.rows[3].cells[0].text = 'Average Order Value'
        stats_table.rows[3].cells[1].text = f"${overall_stats.get('average_order_value', 0):,.2f}"
        stats_table.rows[4].cells[0].text = 'Unique Customers'
        stats_table.rows[4].cells[1].text = str(overall_stats.get('unique_customers', 0))
        stats_table.rows[5].cells[0].text = 'Delivered Orders'
        stats_table.rows[5].cells[1].text = str(overall_stats.get('delivered_orders', 0))
        stats_table.rows[6].cells[0].text = 'Delivered Revenue'
        stats_table.rows[6].cells[1].text = f"${overall_stats.get('delivered_revenue', 0):,.2f}"
    
    doc.add_paragraph()
    
    # Sales Data by Period
    if sales_data:
        doc.add_heading(f'Sales Data by {period.capitalize()}', level=1)
        sales_table = doc.add_table(rows=min(len(sales_data) + 1, 21), cols=5)
        sales_table.style = 'Light Grid Accent 1'
        
        # Headers
        headers = sales_table.rows[0].cells
        headers[0].text = 'Period'
        headers[1].text = 'Orders'
        headers[2].text = 'Revenue'
        headers[3].text = 'Avg Order'
        headers[4].text = 'Customers'
        
        # Data rows (limit to 20)
        for idx, item in enumerate(sales_data[:20], 1):
            cells = sales_table.rows[idx].cells
            cells[0].text = str(item.get('period_label', ''))
            cells[1].text = str(item.get('total_orders', 0))
            cells[2].text = f"${item.get('total_revenue', 0):,.2f}"
            cells[3].text = f"${item.get('average_order_value', 0):,.2f}"
            cells[4].text = str(item.get('unique_customers', 0))
    
    doc.add_paragraph()
    
    # Top Products
    if top_products:
        doc.add_heading('Top 10 Selling Products', level=1)
        products_table = doc.add_table(rows=len(top_products) + 1, cols=4)
        products_table.style = 'Light Grid Accent 1'
        
        # Headers
        headers = products_table.rows[0].cells
        headers[0].text = 'Rank'
        headers[1].text = 'Product Name'
        headers[2].text = 'Quantity Sold'
        headers[3].text = 'Revenue'
        
        # Data rows
        for idx, product in enumerate(top_products, 1):
            cells = products_table.rows[idx].cells
            cells[0].text = str(idx)
            cells[1].text = str(product.get('name', ''))
            cells[2].text = str(product.get('total_quantity_sold', 0))
            cells[3].text = f"${product.get('total_revenue', 0):,.2f}"
    
    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ========== ADMIN USER MANAGEMENT ENDPOINTS ==========
@admin_bp.route('/users', methods=['GET'])
def get_all_users():
    """Get all users (admin view)"""
    try:
        conn = get_db()
        if not conn:
            return jsonify({'error': 'Database connection failed'}), 500
        
        cursor = conn.cursor(dictionary=True)
        
        # Get all users with order statistics
        cursor.execute("""
            SELECT 
                u.id, u.username, u.email, u.first_name, u.last_name, 
                u.phone, u.address, u.is_active, u.is_admin, u.created_at,
                COUNT(DISTINCT o.id) as total_orders,
                SUM(CASE WHEN o.status = 'delivered' THEN o.total_amount ELSE 0 END) as total_spent
            FROM users u
            LEFT JOIN orders o ON u.id = o.user_id
            GROUP BY u.id
            ORDER BY u.created_at DESC
        """)
        
        users = cursor.fetchall()
        
        cursor.close()
        conn.close()
        
        # Convert Decimal to float
        for user in users:
            if 'total_spent' in user and user['total_spent'] is not None:
                user['total_spent'] = float(user['total_spent'])
        
        return jsonify(users), 200
        
    except Exception as e:
        print(f"❌ Get all users error: {e}")
        return jsonify({'error': str(e)}), 500

@admin_bp.route('/users/<int:user_id>/reset-password', methods=['PUT'])
def reset_user_password(user_id):
    """Reset a user's password (admin action)"""
    try:
        data = request.get_json()
        new_password = data.get('new_password')
        
        if not new_password or len(new_password) < 6:
            return jsonify({'error': 'Password must be at least 6 characters'}), 400
        
        conn = get_db()
        if not conn:
            return jsonify({'error': 'Database connection failed'}), 500
        
        cursor = conn.cursor(dictionary=True)
        
        # Check if user exists
        cursor.execute("SELECT id, username, email FROM users WHERE id = %s", (user_id,))
        user = cursor.fetchone()
        
        if not user:
            cursor.close()
            conn.close()
            return jsonify({'error': 'User not found'}), 404
        
        # Hash the new password
        hashed = bcrypt.hashpw(new_password.encode('utf-8'), bcrypt.gensalt())
        
        # Update password
        cursor.execute("""
            UPDATE users 
            SET password_hash = %s, updated_at = NOW()
            WHERE id = %s
        """, (hashed.decode('utf-8'), user_id))
        
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({
            'message': 'Password reset successfully',
            'user': {
                'id': user['id'],
                'username': user['username'],
                'email': user['email']
            }
        }), 200
        
    except Exception as e:
        print(f"❌ Reset password error: {e}")
        return jsonify({'error': str(e)}), 500

@admin_bp.route('/users/<int:user_id>/deactivate', methods=['PUT'])
def deactivate_user(user_id):
    """Deactivate a user account (admin action)"""
    try:
        conn = get_db()
        if not conn:
            return jsonify({'error': 'Database connection failed'}), 500
        
        cursor = conn.cursor(dictionary=True)
        
        # Check if user exists
        cursor.execute("SELECT id, username, email, is_active FROM users WHERE id = %s", (user_id,))
        user = cursor.fetchone()
        
        if not user:
            cursor.close()
            conn.close()
            return jsonify({'error': 'User not found'}), 404
        
        if not user['is_active']:
            cursor.close()
            conn.close()
            return jsonify({'message': 'User is already deactivated'}), 200
        
        # Deactivate user
        cursor.execute("""
            UPDATE users 
            SET is_active = 0, updated_at = NOW()
            WHERE id = %s
        """, (user_id,))
        
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({
            'message': 'User deactivated successfully',
            'user': {
                'id': user['id'],
                'username': user['username'],
                'email': user['email'],
                'is_active': False
            }
        }), 200
        
    except Exception as e:
        print(f"❌ Deactivate user error: {e}")
        return jsonify({'error': str(e)}), 500

@admin_bp.route('/users/<int:user_id>/activate', methods=['PUT'])
def activate_user(user_id):
    """Activate a user account (admin action)"""
    try:
        conn = get_db()
        if not conn:
            return jsonify({'error': 'Database connection failed'}), 500
        
        cursor = conn.cursor(dictionary=True)
        
        # Check if user exists
        cursor.execute("SELECT id, username, email, is_active FROM users WHERE id = %s", (user_id,))
        user = cursor.fetchone()
        
        if not user:
            cursor.close()
            conn.close()
            return jsonify({'error': 'User not found'}), 404
        
        if user['is_active']:
            cursor.close()
            conn.close()
            return jsonify({'message': 'User is already active'}), 200
        
        # Activate user
        cursor.execute("""
            UPDATE users 
            SET is_active = 1, updated_at = NOW()
            WHERE id = %s
        """, (user_id,))
        
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({
            'message': 'User activated successfully',
            'user': {
                'id': user['id'],
                'username': user['username'],
                'email': user['email'],
                'is_active': True
            }
        }), 200
        
    except Exception as e:
        print(f"❌ Activate user error: {e}")
        return jsonify({'error': str(e)}), 500

# ========== PRODUCT IMAGE UPLOAD ENDPOINT ==========
@admin_bp.route('/products/<int:product_id>/upload-image', methods=['POST'])
def upload_product_image(product_id):
    """Upload product image file"""
    try:
        if 'image_file' not in request.files:
            return jsonify({'error': 'No image file provided'}), 400
        
        file = request.files['image_file']
        
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        # Validate file extension
        allowed_extensions = {'png', 'jpg', 'jpeg', 'gif'}
        filename = secure_filename(file.filename)
        file_ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
        
        if file_ext not in allowed_extensions:
            return jsonify({'error': 'Invalid file type. Allowed: png, jpg, jpeg, gif'}), 400
        
        # Create upload directory if it doesn't exist (using constant defined at top of file)
        os.makedirs(UPLOAD_DIR, exist_ok=True)
        
        # Generate unique filename
        unique_filename = f"{uuid.uuid4().hex}.{file_ext}"
        file_path = os.path.join(UPLOAD_DIR, unique_filename)
        
        # Save file
        file.save(file_path)
        
        # Generate absolute URL path using request host
        base_url = request.host_url.rstrip('/')
        image_url = f"{base_url}/static/uploads/products/{unique_filename}"
        
        # Update product in database
        conn = get_db()
        if not conn:
            return jsonify({'error': 'Database connection failed'}), 500
        
        cursor = conn.cursor()
        cursor.execute("UPDATE products SET image_url = %s WHERE id = %s", (image_url, product_id))
        
        if cursor.rowcount == 0:
            cursor.close()
            conn.close()
            # Clean up uploaded file if product not found
            if os.path.exists(file_path):
                os.remove(file_path)
            return jsonify({'error': 'Product not found'}), 404
        
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({
            'message': 'Image uploaded successfully',
            'image_url': image_url
        }), 200
        
    except Exception as e:
        print(f"❌ Upload product image error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500
